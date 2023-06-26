from docx import Document
from docx.shared import Pt, Mm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Optional, Union, List
from xmlrpc.client import Boolean
from collections import OrderedDict
from math import ceil, floor
import re
import numpy as np
import pandas as pd
import os
import subprocess

def mergeDict(dict1, dict2) -> dict:
    if isinstance(dict1, dict) and isinstance(dict2, dict):
        return {
            **dict1, **dict2,
            **{k: mergeDict(dict1[k], dict2[k])
            for k in {*dict1} & {*dict2}}
        }
    else:
        return [
            *(dict1 if isinstance(dict1, list) else [dict1]),
            *(dict2 if isinstance(dict2, list) else [dict2])
        ]

class ExcelParser:
    def __init__(
                self,
                scan_path: Optional[Union[str, None]] = None,
                export_path: Optional[Union[str, None]] = None,
                server_types_filter: Optional[Union[list, None]] = None,
                countours_filter: Optional[Union[list, None]] = None,
                nodes_filter: Optional[Union[list, None]] = None
                ) -> None:
        self._convertor_formats = ['.xls', '.xlsx', '.xlsm']
        self._filters = []

        self._formated_table = []
        self._excel_df = pd.DataFrame
        self._excel_df_file = ""

        if server_types_filter is None:
            self.server_types = []
        else:
            self.server_types = server_types_filter

        if countours_filter is None:
            self.countours = []
        else:
            self.countours = countours_filter

        if nodes_filter is None:
            self.nodes = []
        else:
            self.nodes = nodes_filter 

        if export_path is None:
            self.export_path = self._get_current_workdir()
        else:
            self.export_path = export_path

        if scan_path is None:
            self.scan_path = self._get_current_workdir()
        else:
            self.scan_path = scan_path

    def _formated_table2formated_dict(self) -> Union[dict, None]:
        list_len = len(self._formated_table)
        if list_len > 2:
            merged_dict = self._formated_table[0]
            for table_item in self._formated_table[1:]:
                merged_dict = mergeDict(table_item, merged_dict)
            return merged_dict
        elif list_len == 1:
            return self._formated_table[0]
        else:
            return None

    def _set_excel_file(self, path : str) -> Boolean:
        _, ext = os.path.splitext(path)
        if ext in self._convertor_formats:
            self._excel_df = pd.read_excel(path)
            return True
        else:
            return False

    def _get_filters(self) -> list:
        self._filters.clear()

        if self.countours:
            self._filters.append(self._excel_df["Тип"].str.contains(
                '|'.join(self.countours),
                case=False,
                na=False).values
                )

        if self.nodes:
            self._filters.append(self._excel_df["Имя узла"].isin(self.nodes))

        if self.server_types:
            row_server_types = []
            for _, df_row in self._excel_df.iterrows():
                for server_type in self.server_types:
                    if server_type in str(df_row["Описание"]):
                        row_server_types.append(server_type)
                    else:
                        row_server_types.append("")

            self._filters.append(pd.Series(row_server_types, name="Тип_машины").str.contains(
                '|'.join(self.server_types),
                case=False,
                na=False).values
                )

    def _filter_dataframe(self) -> pd.DataFrame:
        self._get_filters() 

        if self._filters:
            return self._excel_df.loc[np.logical_and.reduce(self._filters)]
        else:
            return self._excel_df

    def _get_current_workdir(self) -> str:
        cmd_result = subprocess.Popen(["pwd"], stdout=subprocess.PIPE)
        return cmd_result.communicate()[0].decode('utf-8').strip("\n")

    def get_scandir_files(self) -> tuple:
        error = False
        if os.path.exists(self.scan_path):
            cmd_result = subprocess.Popen(["ls", self.scan_path], stdout=subprocess.PIPE)
            return error, ["{}/{}".format(self.scan_path, file) for file in cmd_result.communicate()[0].decode('utf-8').strip("\n").split("\n")]
        error =True
        return error, None

    def _get_hostname_from_excel(self,
                formated_table_data : dict,
                excel_row : pd.Series
                ) -> tuple:
        error = False
        hostname_pattern = ".*"
        if re.search(hostname_pattern, str(excel_row["Хостнейм"])):
            formated_table_data["hostname"] = str(excel_row["Хостнейм"])
        else:
            formated_table_data.clear()
            error = True

        return error, formated_table_data

    def _get_region_from_excel(
                self,
                formated_table_data : dict,
                excel_row : pd.Series
                ) -> tuple:
        error = False
        if str(excel_row["Имя узла"]):
            formated_table_data["region"] = str(excel_row["Имя узла"])
        else:
            formated_table_data.clear()
            error = True

        return error, formated_table_data

    def _get_cidr_from_excel(
                self,
                formated_table_data : dict,
                excel_row : pd.Series
                ) -> tuple:
        error = False
        cidr_pattern = "^([01]?\d\d?|2[0-4]\d|25[0-5])(?:\.(?:[01]?\d\d?|2[0-4]\d|25[0-5])){3}(?:/[0-2]?\d|/3[0-2])?$"
        if re.search(cidr_pattern, str(excel_row["Подсеть"])):
            formated_table_data["network"] = str(excel_row["Подсеть"])
        else:
            formated_table_data.clear()
            error = True

        return error, formated_table_data

    def format_excel_table(self, file_path: str) -> Union[pd.DataFrame, None]:
        self._excel_df_file = file_path
        if not self._set_excel_file(self._excel_df_file):
            return None

        table_df = self._filter_dataframe()

        for _, excel_row in table_df.iterrows():
            formated_table_data = {
                "hostname": None,
                "network": None,
                "region": None,
            }

            error, formated_table_data = self._get_hostname_from_excel(formated_table_data, excel_row)
            if error:
                continue
            
            error, formated_table_data = self._get_cidr_from_excel(formated_table_data, excel_row)
            if error:
                continue

            error, formated_table_data = self._get_region_from_excel(formated_table_data, excel_row)
            if error:
                continue
            
            self._formated_table.append(formated_table_data)

        if self._formated_table:
            formated_dict = self._formated_table2formated_dict()
            if formated_dict is not None:
                frame = pd.DataFrame.from_dict(formated_dict, orient='index').transpose()
            self._formated_table.clear()
            return frame
        else:
            return None

    def merge_tables(self, df_list: List[pd.DataFrame]) -> pd.DataFrame:
        return pd.concat(df_list, ignore_index=True)

    def run(self) -> None:
        np.warnings.filterwarnings('ignore', category=np.VisibleDeprecationWarning) 

        error, files_list = self.get_scandir_files()
        if error:
            print("Неверная папка для сканирования")
        else:
            list_tables_dataframes = []
            for file in files_list:
                new_df = self.format_excel_table(file)
                if new_df is not None:
                    list_tables_dataframes.append(new_df)
            merged_table = self.merge_tables(list_tables_dataframes)

        return merged_table


class WordPageContent:

        class WordPageView:
            pass

        class SideLabelsTableView(WordPageView):
            def __init__(self, rows : int, cols : int) -> None:
                self.table_rows = rows
                self.table_cols = cols
                self.table_cells_side_label = []

            def add_cell(self, cell : dict):
                self.table_cells_side_label.append(cell)

        class FrontsideLabelsTableView(WordPageView):
            def __init__(self, rows : int, cols : int) -> None:
                self.table_rows = rows
                self.table_cols = cols
                self.table_cells_frontside_label = []
            
            def add_cell(self, cell : dict):
                self.table_cells_frontside_label.append(cell)

        def __init__(self) -> None:
            self.page_objects = []

        def add_page_object(self, obj : WordPageView):
            self.page_objects.append(obj)

class WordPage:

        class PageObject:
            def __init__(self, name = None, ref = None) -> None:
                self.object_name = name
                self.object_ref = ref

        def __init__(self) -> None:
            self.page_objects = []

        def add_object(self, object : PageObject):
            self.page_objects.append(object)

        def get_object_by_name(self, name : str):
            for object in self.page_objects:
                if object.object_name == name:
                    return object

class AbstractDataStorage:
    def __init__(self) -> None:
        self.side_label_data = []
        self.frontside_label_data = []

    def get_side_label_item_by_idx(self, idx : int):
        if idx >= len(self.side_label_data):
            return SideLabelData()
        else:
            return self.side_label_data[idx]

    def get_frontside_label_item_by_idx(self, idx : int):
        if idx >= len(self.frontside_label_data):
            return FrontsideLabelData()
        else:
            return self.frontside_label_data[idx]
            
        
class AbstractData:
    def __init__(self, hostname="", ip="", project_number="") -> None:
        self.data_struct = OrderedDict()
        self.data_struct["Hostname"] = hostname
        self.data_struct["IP"] = ip
        self.data_struct["Проектный номер"] = str(project_number)

    def get_struct(self):
        return self.data_struct


class SideLabelData(AbstractData):
        def __init__(self, zone="", weight="", *args, **kwargs) -> None:
            super(SideLabelData, self).__init__(*args, **kwargs)
            self.data_struct["Площадка"] = zone
            self.data_struct["Вес"] = str(weight)


class FrontsideLabelData(AbstractData):
        def __init__(self, *args, **kwargs) -> None:
            super(FrontsideLabelData, self).__init__(*args, **kwargs)


class WordLabelGenerator:

    class RawDataExtractor:
        def __init__(self, df: pd.DataFrame, repeate_count: int) -> None:
            self.dataframe = df
            self.repeate_count = repeate_count

        def _get_network(self, row : pd.Series) -> str:
            return str(row["network"])

        def _get_hostname(self, row : pd.Series) -> str:
            return str(row["hostname"])

        def _get_region(self, row : pd.Series) -> str:
            return str(row["region"])

        def run(self) -> AbstractDataStorage:
            data_storage = AbstractDataStorage()
            for _, row in self.dataframe.iterrows():
                hostname = self._get_hostname(row)
                network = self._get_network(row)
                zone = self._get_region(row)

                for _ in range(self.repeate_count):
                    data_storage.frontside_label_data.append(FrontsideLabelData(
                        hostname=hostname, ip=network, project_number=1))
                    data_storage.side_label_data.append(SideLabelData(
                        hostname=hostname, ip=network,
                        zone=zone, project_number=1, weight="25 кг."))
            else:
                return data_storage

    def __init__(self, input_df : pd.DataFrame, repeate_count: int) -> None:
        if repeate_count >= 1:
            self.repeate_count = repeate_count
        else:
            self.repeate_count = 1

        self.word_doc = Document(docx=None)
        self.word_view_pages = self._get_word_view_pages(input_df)
        self.word_page = WordPage()

    def _get_word_view_pages(self, df : pd.DataFrame) -> List:
        word_pages_view = []
        extractor = self.RawDataExtractor(df, self.repeate_count)
        data_storage = extractor.run()
 
        frontside_cells_count = 14 * 4
        page_count_frontside = ceil(len(data_storage.frontside_label_data) / frontside_cells_count)
        frontside_label_item_idx = 0
        new_view = WordPageContent()
        new_view.add_page_object(WordPageContent.FrontsideLabelsTableView(page_count_frontside * 14, 4))

        for _ in range(new_view.page_objects[0].table_rows):
            for _ in range(new_view.page_objects[0].table_cols):
                new_view.page_objects[0].add_cell(data_storage.get_frontside_label_item_by_idx(frontside_label_item_idx))
                frontside_label_item_idx += 1
        else:
            word_pages_view.append(new_view)

        side_cells_count = 7 * 2
        page_count_side = ceil(len(data_storage.side_label_data) / side_cells_count)
        side_label_item_idx = 0
        new_view = WordPageContent()
        new_view.add_page_object(WordPageContent.SideLabelsTableView(page_count_side * 7, 2))

        for _ in range(new_view.page_objects[0].table_rows):
            for _ in range(new_view.page_objects[0].table_cols):
                new_view.page_objects[0].add_cell(data_storage.get_side_label_item_by_idx(side_label_item_idx))
                side_label_item_idx += 1
        else:
            word_pages_view.append(new_view)

        return word_pages_view

    def _set_doc_format(self):
        section = self.word_doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(10)
        section.right_margin = Mm(7.5)
        section.top_margin = Mm(7.5)
        section.bottom_margin = Mm(0)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)

    def _create_table(self, name, rows, cols):
        table = self.word_doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.autofit = False
        new_object = WordPage.PageObject(name, table)
        self.word_page.add_object(new_object)

    def _add_table_content(self, table_counter : int, label_data : dict, tab_name : str):
        table = self.word_page.get_object_by_name(tab_name).object_ref
        cells_length_row = len(table.rows[0].cells)
        row_number = floor(table_counter / cells_length_row)
        cell_number = table_counter % cells_length_row
        row = table.rows[row_number]
        cell = row.cells[cell_number]
        cell_data = ["{}: {}".format(key, value) for key, value in label_data.items()]
        cell.text = "\n".join(cell_data)

    def _add_frontside_table_style(self, table):
        for row in table.rows:
            row.height = Mm(20)
            for cell in row.cells:
                cell.width = Mm(48.5)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell_paragraph = cell.paragraphs[0].runs[0]
                cell_paragraph.font.bold = True
                cell_paragraph.font.size = Pt(10)
                cell_paragraph.font.name = "Times New Roman"

    def _add_side_table_style(self, table):
        for row in table.rows:
            row.height = Mm(40)
            for cell in row.cells:
                cell.width = Mm(97)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell_paragraph = cell.paragraphs[0].runs[0]
                cell_paragraph.font.bold = True
                cell_paragraph.font.size = Pt(14)
                cell_paragraph.font.name = "Times New Roman"
                cell.paragraphs[0].alignment= WD_PARAGRAPH_ALIGNMENT.CENTER

    def _init_side_table(self, page_view_side):
        table_counter = 0
        self._create_table("side_label_table", page_view_side.page_objects[0].table_rows,
                                                        page_view_side.page_objects[0].table_cols)
        for side_cell_content in page_view_side.page_objects[0].table_cells_side_label:
            self._add_table_content(table_counter, side_cell_content.get_struct(), "side_label_table")
            table_counter += 1
        else:
            self._add_side_table_style(self.word_page.get_object_by_name("side_label_table").object_ref)

    def _init_frontside_table(self, page_view_frontside):
        table_counter = 0
        self._create_table("frontside_label_table", page_view_frontside.page_objects[0].table_rows,
                                                    page_view_frontside.page_objects[0].table_cols)
        for frontside_cell_content in page_view_frontside.page_objects[0].table_cells_frontside_label:
            self._add_table_content(table_counter, frontside_cell_content.get_struct(), "frontside_label_table")
            table_counter += 1
        else:
            self._add_frontside_table_style(self.word_page.get_object_by_name("frontside_label_table").object_ref)

    def _create_word_template(self):
        self._set_doc_format()

        page_view_frontside = None
        for page_view in self.word_view_pages:
            if not isinstance(page_view.page_objects[0], WordPageContent.FrontsideLabelsTableView):
                continue 
            page_view_frontside = page_view
            break

        self._init_frontside_table(page_view_frontside)

        page_view_side = None
        for page_view in self.word_view_pages:
            if not isinstance(page_view.page_objects[0], WordPageContent.SideLabelsTableView):
                continue
            page_view_side = page_view
            break

        self._init_side_table(page_view_side)


    def excel2Word_label(self):
        self._create_word_template() 
        self.word_doc.save('labels.docx')

if __name__ == "__main__":
    excel_parser = ExcelParser(
        scan_path="./excel.d/"
    )

    filtered_df = excel_parser.run()
    generator = WordLabelGenerator(filtered_df, 2)
    generator.excel2Word_label()